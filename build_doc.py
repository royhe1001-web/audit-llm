#!/usr/bin/env python3
"""
build_doc.py — 通用 Markdown → Word 转换器
==========================================
输入: MD 文件路径
输出: Word 文档
样式:
  - 字体:宋体
  - 首行缩进:两字符
  - 行距:1.5 倍
  - 分级标题:H1/H2/H3
  - 标点:全角(数字内小数点保留半角)

支持语法:
  - # 标题 → 一级标题(15pt 加粗)
  - ## 标题 → 二级标题(13pt 加粗)
  - ### 标题 → 三级标题(12pt 加粗)
  - #### 标题 → 四级标题(11pt 加粗)
  - 普通段落 → 正文(12pt,首行缩进 0.74cm)
  - | 表头 | → 表格(深蓝表头)
  - ---  → 段落分隔(空行)
  - > 引用 → 引用块(左缩进,11pt)
  - ** 加粗 ** → 加粗
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============ 样式函数 ============
LQUO = '“'
RQUO = '”'
EM_DASH = '—'


def to_full_punct(s):
    """中文文本 ASCII 标点转全角(数字内小数点/英文字母缩写保留半角)"""
    if not s:
        return s
    # 保护小数点:数字.数字
    s = re.sub(r'(\d)\.(\d)', lambda m: m.group(1) + '\x00' + m.group(2), s)
    # 保护英文缩写:字母.字母 (如 e.g., i.e., U.S.)
    s = re.sub(r'([a-zA-Z])\.([a-zA-Z])', lambda m: m.group(1) + '\x01' + m.group(2), s)
    # 保护转义双引号(在 add_para 内)
    s = s.replace('\\"', '\x02')
    # 双引号成对替换为左右全角引号
    parts = s.split('"')
    if len(parts) >= 3:
        rebuilt = parts[0]
        is_open = True
        for p in parts[1:]:
            rebuilt += (LQUO if is_open else RQUO) + p
            is_open = not is_open
        s = rebuilt
    # ASCII 标点转全角(此时 . 已不会被数字/字母破坏)
    s = s.replace('.', '。')
    s = s.replace(',', '，')
    s = s.replace(';', '；')
    s = s.replace(':', '：')
    s = s.replace('?', '？')
    s = s.replace('!', '！')
    s = s.replace('(', '（')
    s = s.replace(')', '）')
    s = s.replace('...', '……')
    # 还原小数点和缩写
    s = s.replace('\x00', '.')
    s = s.replace('\x01', '.')
    s = s.replace('\x02', '"')
    return s


def set_zh_font(run, size_pt=12, bold=False, color=(0, 0, 0)):
    if run.text:
        run.text = to_full_punct(run.text)
    run.font.name = '宋体'
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_zh_font(run, size_pt=20, bold=True)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_zh_font(run, size_pt=15, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_zh_font(run, size_pt=13, bold=True)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_zh_font(run, size_pt=12, bold=True)


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 4']
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_zh_font(run, size_pt=11, bold=True)


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_zh_font(run, size_pt=12)


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_zh_font(run, size_pt=11, color=(80, 80, 80))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_zh_font(run, size_pt=10.5, bold=True, color=(255, 255, 255))
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '185FA5')
        tcPr.append(shd)
    # 数据行
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = ''
            para = row_cells[c_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(val))
            set_zh_font(run, size_pt=10.5)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('———')
    set_zh_font(run, size_pt=12, color=(150, 150, 150))


def add_hr(doc):
    """水平分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('—' * 20)
    set_zh_font(run, size_pt=10, color=(180, 180, 180))


# ============ 表格解析 ============
def parse_table(lines, idx):
    """从 lines[idx] 开始解析 markdown 表格,返回 (headers, rows, next_idx)"""
    table_lines = []
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx])
        idx += 1
    if len(table_lines) < 3:
        return None, None, idx
    # 解析表头
    headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
    # 跳过分隔行(table_lines[1])
    # 解析数据行
    rows = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return headers, rows, idx


# ============ 主解析函数 ============
def md_to_word(md_path, doc_path):
    md_text = Path(md_path).read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                # 输出代码块(简化处理:作为一个段落)
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.74)
                    p.paragraph_format.line_spacing = 1.15
                    run = p.add_run('\n'.join(code_lines))
                    set_zh_font(run, size_pt=10)
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 标题
        if stripped.startswith('# '):
            add_title(doc, stripped[2:].strip())
        elif stripped.startswith('## '):
            add_h1(doc, stripped[3:].strip())
        elif stripped.startswith('### '):
            add_h2(doc, stripped[4:].strip())
        elif stripped.startswith('#### '):
            add_h3(doc, stripped[5:].strip())
        elif stripped.startswith('##### '):
            add_h4(doc, stripped[6:].strip())
        # 分隔符(讲稿用)— 输出空段不画水平线
        elif stripped == '---' or stripped == '———':
            # 讲稿风格:跳过水平线,只用空行分隔
            # add_hr(doc)  # 注释掉水平线
            doc.add_paragraph()  # 输出一个空行
            i += 1
            continue
        # 引用块
        elif stripped.startswith('> '):
            add_quote(doc, stripped[2:].strip())
        # 表格
        elif stripped.startswith('|'):
            headers, rows, next_i = parse_table(lines, i)
            if headers is not None:
                add_table(doc, headers, rows)
                i = next_i
                continue
        # 普通段落
        else:
            txt = stripped
            # 处理粗体 **xxx** → 用 set_zh_font bold=True 处理(简化:这里直接显示)
            txt = re.sub(r'\*\*(.+?)\*\*', r'\1', txt)
            txt = re.sub(r'\*(.+?)\*', r'\1', txt)
            txt = re.sub(r'`(.+?)`', r'\1', txt)
            # 去掉块级 markdown 链接 [text](url) → text
            txt = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', txt)
            add_para(doc, txt)
        i += 1

    doc.save(doc_path)
    print(f"✅ Word 已生成: {doc_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python build_doc.py <md_path> <docx_path>")
        sys.exit(1)
    md_to_word(sys.argv[1], sys.argv[2])